from docx import Document
import os
import json

INPUT_FILE = 'annotated_docx'  # r'C:\Users\Super-IdoI\Desktop\辅导实习\NER辅导项目\标注'
OUTPUT_DIR = 'dataset_final'  # r'C:\Users\Super-IdoI\Desktop\辅导实习\NER辅导项目\BERT_LSTM_CRF\data\dataset_final'
RATIO = 0.95/0.05  # = 19 : 1
SPLIT_TOKEN = ['。', '，', '；']
train_and_test = 'train_and_test_pre.json'
train = 'train_pre.json'
test = 'test_pre.json'

"""
将文件夹 INPUT_FILE 下的所有docx文件转换为下面格式的标注数据
{"text": "彭小军认为，国内银行现在走的是台湾的发卡模式，先通过跑马圈地再在圈的地里面选择客户，", "label": {"address": {"台湾": [[15, 16]]}, "name": {"彭小军": [[0, 2]]}}}
存入 文件夹 INPUT_FILE ，并 按照 RATIO 拆分训练、测试数据集 
"""


def preprocessing_func():
    if not os.path.exists(os.path.join(OUTPUT_DIR,'source_data.json')):
        if not os.path.exists(INPUT_FILE):
            print(f'没有{INPUT_FILE}文件夹.')
            exit(-1)
        if not os.path.exists(OUTPUT_DIR):
            os.mkdir(OUTPUT_DIR)
        source_dict = {}
        for filename in os.listdir(INPUT_FILE):
            filepath = os.path.join(INPUT_FILE,filename)
            file_content = ''
            redText = []
            if filepath.endswith('.docx'):
                docx_f = Document(filepath)
                for para in docx_f.paragraphs:
                    # print(para)
                    file_content += para.text
                    for r in para.runs:
                        try:
                            if r.font.highlight_color == 6:
                                redText.append(r.text)
                        except Exception as e:
                            print(e.args)
                            print(f'{filename}: {r.text} 处理中遇到一处错误.')
                if len(redText) == 0:
                    print(f'{filepath}： 由于某些原因，未检测到标注的金融领域实体.')
                else:
                    source_dict[filename] = (file_content, redText)
        with open(os.path.join(OUTPUT_DIR,'source_data.json'),mode='wt',encoding='utf-8') as outp:
            json.dump(source_dict, outp,ensure_ascii=False,indent=4)

    with open(os.path.join(OUTPUT_DIR,'source_data.json'),mode='rt',encoding='utf-8') as outp:
        source_dict = json.load(outp)

    samples = []
    out_tr_te = os.path.join(OUTPUT_DIR,train_and_test)
    if not os.path.exists(out_tr_te):
        outp = open(out_tr_te, mode='wt',encoding='utf-8')
        for filename, (contents, entities) in source_dict.items():
            entities = list(set(entities))
            # print(contents, entities)

            # 寻找文章中所有出现大于一次的金融实体
            morethanone_entities = []
            for entity in entities:
                if contents.count(entity)>1:
                    morethanone_entities.append(entity)
            print(f'{filename} 中有 {len(morethanone_entities)} 个出现大于一次的金融实体')
            has_write_entity = []
            for entity in entities:
                if entity not in has_write_entity: #True: # if contents.count(entity) < 2:
                    start_idx = contents.find(entity)
                    if start_idx != -1:
                        # （寻找句子开始位置） 前向遍历，碰到第一个分隔符且至少5个字符才结束.
                        sent_start_idx = 0
                        for idx in range(start_idx,0,-1):
                            if (contents[idx] in SPLIT_TOKEN and start_idx- idx >= 5) or start_idx- idx >= 50:
                                if idx !=0:
                                    sent_start_idx = idx +1
                                break
                        # （寻找句子结束位置） 后向遍历，碰到第一个分隔符且至少10个字符
                        sent_end_idx = len(contents)-1
                        for idx in range(start_idx, len(contents),1):
                            if (contents[idx] in SPLIT_TOKEN and idx - start_idx >= 10) or idx - start_idx >= 50:
                                sent_end_idx = idx
                                break
                        sentence = contents[sent_start_idx: sent_end_idx] # "彭小军认为，国内银行现在走的是台湾的发卡模式，先通过跑马圈地再在圈的地里面选择客户，"
                        entities_info = {}  # {"台湾": [[15, 16]]，"中国": [[13, 14]]}
                        for entity_ in entities:
                            entity_idx = sentence.find(entity_)
                            while entity_idx != -1:
                                if entity_ not in has_write_entity and entity_ not in morethanone_entities:
                                    has_write_entity.append(entity_)
                                if entity_ not in entities_info:
                                    entities_info[entity_] = [[entity_idx, entity_idx+len(entity_)-1]]
                                else:
                                    entities_info[entity_].append([entity_idx, entity_idx+len(entity_)-1])
                                entity_idx = sentence.find(entity_, entity_idx+1)
                        sample = {'text':sentence, 'label':{'financial_entity':entities_info}}
                        samples.append(sample)
                        outp.write(json.dumps(sample,ensure_ascii=False)+'\n')
        outp.close()
    else:  # 已经存在 'train_and_test_pre.json'， 直接读取即可
        with open(out_tr_te, mode='rt',encoding='utf-8') as inp:
            for line in inp.readlines():
                sample = eval(line.strip())
                samples.append(sample)
    # 目前的samples对象便可切分了
    import random
    random.shuffle(samples)
    count_train = 0
    with open(os.path.join(OUTPUT_DIR,train), mode='wt',encoding='utf-8') as outp:
        for sample in samples[:int(len(samples)*RATIO/(RATIO+1))]:
            count_train +=1
            outp.write(json.dumps(sample,ensure_ascii=False)+'\n')
    print(f'写入{count_train} 至 {train}')
    count_test = 0
    with open(os.path.join(OUTPUT_DIR,test), mode='wt', encoding='utf-8') as outp:
        for sample in samples[int(len(samples)*RATIO/(RATIO+1)):]:
            count_test +=1
            outp.write(json.dumps(sample,ensure_ascii=False)+'\n')
    print(f'写入{count_test} 至 {test}')


if __name__ == '__main__':
    TEST_FLAG = False
    if TEST_FLAG:
        redText = []
        doc = Document('./万科.docx')

        for p in doc.paragraphs:
            for r in p.runs:
                if r.font.highlight_color == 6:
                    redText.append(r.text)
        print(redText)
    else:
        preprocessing_func()